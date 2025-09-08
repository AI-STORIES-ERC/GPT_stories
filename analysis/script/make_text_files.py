import pandas as pd
import os
from docx import Document

def make_text_files(country, num_stories):
    """
    Create .txt and .docx files from a CSV file containing stories.
    
    Parameters:
    - country: The country code for the file (e.g. "AU").
    - num_stories: The number of stories to process.
    """
    input_path = f"data/{country}/{country}_stories.csv"
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"The file {input_path} does not exist.")
    
    df = pd.read_csv(input_path)

    if num_stories > len(df):
        raise ValueError(f"Requested {num_stories} stories, but only {len(df)} available.")
    
    df = df.head(num_stories)

    output_dir = "analysis/data/story_texts"
    os.makedirs(output_dir, exist_ok=True)

    for row in df.itertuples(index=False):
        story_id = row.Story_ID
        story_text = row.Story

        # Save as text file
        txt_path = os.path.join(output_dir, f"{story_id}.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(story_text)

        # Save as Word document
        doc = Document()
        doc.add_heading(story_id, level=1)
        doc.add_paragraph(story_text)
        doc_path = os.path.join(output_dir, f"{story_id}.docx")
        doc.save(doc_path)

    print(f"{num_stories} stories saved to {output_dir} as .txt and .docx")

if __name__ == "__main__":
    make_text_files("RU", 50)
